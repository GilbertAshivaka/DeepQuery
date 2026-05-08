"""
Deep Query — DOCX Parser (python-docx)

Extracts text (with heading hierarchy), tables, and images from Word documents.
"""

import io
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from ingestion.parser_pdf import ExtractedBlock

logger = logging.getLogger(__name__)


class DOCXParser:
    """Parse DOCX documents using python-docx.

    Extracts:
    - Paragraph text with heading hierarchy preserved as metadata
    - Embedded tables as structured text
    - Embedded images as separate assets for multimodal embedding
    """

    def parse(self, file_path: Path) -> List[ExtractedBlock]:
        """Parse a DOCX file and return extracted blocks."""
        doc = DocxDocument(str(file_path))
        blocks: List[ExtractedBlock] = []
        current_heading = ""
        current_heading_level = 0
        heading_stack: List[str] = []

        # ── Extract paragraphs ────────────────────────────────
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower()

            # Track heading hierarchy
            if "heading" in style_name:
                level = self._parse_heading_level(style_name)
                # Trim the heading stack to the current level
                heading_stack = heading_stack[: max(0, level - 1)]
                heading_stack.append(text)
                current_heading = " > ".join(heading_stack)
                current_heading_level = level
                continue  # Headings become context, not standalone blocks

            blocks.append(
                ExtractedBlock(
                    text=text,
                    block_type="text",
                    page_number=0,  # DOCX doesn't have page numbers natively
                    heading_context=current_heading,
                    metadata={
                        "paragraph_index": para_idx,
                        "style": para.style.name,
                    },
                )
            )

        # ── Extract tables ────────────────────────────────────
        for table_idx, table in enumerate(doc.tables):
            table_text = self._extract_table(table)
            if table_text and len(table_text) > 30:
                blocks.append(
                    ExtractedBlock(
                        text=table_text,
                        block_type="table",
                        page_number=0,
                        heading_context=current_heading,
                        metadata={"table_index": table_idx, "is_table": True},
                    )
                )

        # ── Extract images ────────────────────────────────────
        image_blocks = self._extract_images(doc)
        blocks.extend(image_blocks)

        logger.info(f"Parsed DOCX: {file_path.name} → {len(blocks)} blocks")
        return blocks

    def _parse_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name (e.g., 'heading 2' → 2)."""
        for part in style_name.split():
            if part.isdigit():
                return int(part)
        return 1

    def _extract_table(self, table) -> str:
        """Convert a docx table to markdown-style text."""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _extract_images(self, doc: DocxDocument) -> List[ExtractedBlock]:
        """Extract embedded images from the DOCX package."""
        blocks = []

        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    img_bytes = image_part.blob
                    content_type = image_part.content_type

                    # Skip very small images (icons, bullets)
                    if len(img_bytes) < 5000:
                        continue

                    # Determine format
                    fmt = "png"
                    if "jpeg" in content_type or "jpg" in content_type:
                        fmt = "jpeg"

                    blocks.append(
                        ExtractedBlock(
                            text="",
                            block_type="image",
                            page_number=0,
                            image_bytes=img_bytes,
                            image_format=fmt,
                            metadata={"relationship_id": rel_id},
                        )
                    )
                except Exception as e:
                    logger.warning(f"Failed to extract image {rel_id}: {e}")

        return blocks
